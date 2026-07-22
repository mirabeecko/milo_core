#!/usr/bin/env python3
"""Build the polished MiLO design DOCX from the canonical Markdown source."""

from __future__ import annotations

import math
import re
import subprocess
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "MILO_SYSTEM_DESIGN.md"
BUILD = ROOT / "build"
ASSETS = ROOT / "assets"
ARTIFACTS = ROOT.parent / "artifacts"

NAVY = "17324D"
TEAL = "0B7A75"
SKY = "DCEEF2"
PALE = "F4F7F9"
CORAL = "E46F54"
INK = "203040"
MUTED = "5C6B75"
WHITE = "FFFFFF"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    return ImageFont.truetype(path, size=size)


def wrap_text(draw: ImageDraw.ImageDraw, value: str, chosen_font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    words = value.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=chosen_font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, subtitle: str, fill: str, outline: str = NAVY) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=26, fill=f"#{fill}", outline=f"#{outline}", width=4)
    title_font = font(34, True)
    sub_font = font(23)
    title_lines = wrap_text(draw, title, title_font, x2 - x1 - 52)
    sub_lines = wrap_text(draw, subtitle, sub_font, x2 - x1 - 52)
    total = len(title_lines) * 44 + len(sub_lines) * 32 + 8
    y = y1 + (y2 - y1 - total) / 2
    for line in title_lines:
        box = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((x1 + x2 - (box[2] - box[0])) / 2, y), line, font=title_font, fill=f"#{INK}")
        y += 44
    y += 8
    for line in sub_lines:
        box = draw.textbbox((0, 0), line, font=sub_font)
        draw.text(((x1 + x2 - (box[2] - box[0])) / 2, y), line, font=sub_font, fill=f"#{MUTED}")
        y += 32


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = TEAL, width: int = 7) -> None:
    sx, sy = start
    ex, ey = end
    draw.line((sx, sy, ex, ey), fill=f"#{color}", width=width)
    angle = math.atan2(ey - sy, ex - sx)
    size = 22
    left = (ex - size * math.cos(angle - math.pi / 6), ey - size * math.sin(angle - math.pi / 6))
    right = (ex - size * math.cos(angle + math.pi / 6), ey - size * math.sin(angle + math.pi / 6))
    draw.polygon((end, left, right), fill=f"#{color}")


def build_architecture_diagram() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    out = ASSETS / "milo-architecture.png"
    image = Image.new("RGB", (2000, 1500), "#F7FAFB")
    draw = ImageDraw.Draw(image)

    owner = (720, 45, 1280, 190)
    hermes = (650, 255, 1350, 425)
    paperclip = (90, 535, 600, 720)
    atlas = (690, 535, 1050, 720)
    forge = (1100, 535, 1460, 720)
    argus = (1510, 535, 1910, 720)
    dashboard = (90, 870, 600, 1045)
    gateway = (1030, 870, 1660, 1045)
    ledger = (1030, 1135, 1660, 1305)
    views = (1030, 1360, 1660, 1480)

    draw_arrow(draw, (1000, owner[3]), (1000, hermes[1]))
    draw_arrow(draw, (760, hermes[3]), (350, paperclip[1]))
    draw_arrow(draw, (860, hermes[3]), (870, atlas[1]))
    draw_arrow(draw, (1070, hermes[3]), (1280, forge[1]))
    draw_arrow(draw, (1240, hermes[3]), (1710, argus[1]))
    draw_arrow(draw, (350, paperclip[3]), (350, dashboard[1]))
    draw_arrow(draw, (1280, forge[3]), (1280, gateway[1]))
    draw_arrow(draw, (1710, argus[3]), (1510, gateway[1]))
    draw_arrow(draw, (1345, gateway[3]), (1345, ledger[1]))
    draw_arrow(draw, (1345, ledger[3]), (1345, views[1]))

    draw_box(draw, owner, "Miroslav · Owner", "cíle, rozhodnutí, schválení", SKY)
    draw_box(draw, hermes, "Hermes · Chief", "jediný hlavní interface a orchestrátor", "D9F1EC", TEAL)
    draw_box(draw, paperclip, "Paperclip", "tickety · runy · audit · náklady", "E8EDF2")
    draw_box(draw, atlas, "Atlas", "analýza a strategie", "FFF4E8", CORAL)
    draw_box(draw, forge, "Forge", "provedení a výstupy", "FFF4E8", CORAL)
    draw_box(draw, argus, "Argus", "nezávislá kontrola", "FFF4E8", CORAL)
    draw_box(draw, dashboard, "Control dashboard", "live realita práce", "E8EDF2")
    draw_box(draw, gateway, "Truth Gateway", "deterministická append-only brána", "D9F1EC", TEAL)
    draw_box(draw, ledger, "Fact Ledger", "Supabase · fakta a provenance", SKY)
    draw_box(draw, views, "Obsidian + MiLO pohledy", "přegenerovatelné projekce", "E8EDF2")

    image.save(out, optimize=True)
    return out


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Strana ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def build_reference_doc(path: Path) -> None:
    # Start from Pandoc's own reference so all DOCX-specific table/list/TOC
    # styles and numbering definitions remain valid, then apply our design.
    path.write_bytes(subprocess.check_output(["pandoc", "--print-default-data-file", "reference.docx"]))
    doc = Document(path)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(19)
    section.right_margin = Mm(17)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name in ("Body Text", "First Paragraph"):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Aptos"
            style.font.size = Pt(9.5)
            style.font.color.rgb = RGBColor.from_string(INK)
            style.paragraph_format.space_after = Pt(5)
            style.paragraph_format.line_spacing = 1.12

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(31)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(90)
    title.paragraph_format.space_after = Pt(16)

    if "Subtitle" in styles:
        subtitle = styles["Subtitle"]
        subtitle.font.name = "Aptos"
        subtitle.font.size = Pt(15)
        subtitle.font.color.rgb = RGBColor.from_string(TEAL)
        subtitle.paragraph_format.space_after = Pt(28)

    heading_settings = {
        "Heading1": (18, NAVY, 18, 8),
        "Heading2": (13, TEAL, 13, 5),
        "Heading3": (10.5, CORAL, 9, 3),
    }
    for style_id, (size, color, before, after) in heading_settings.items():
        style = next(item for item in styles if item.style_id == style_id)
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "TOC Heading" in styles:
        toc = styles["TOC Heading"]
        toc.font.name = "Aptos Display"
        toc.font.size = Pt(20)
        toc.font.bold = True
        toc.font.color.rgb = RGBColor.from_string(NAVY)

    for name in ("Caption", "Image Caption"):
        if name in styles:
            style = styles[name]
            style.font.name = "Aptos"
            style.font.size = Pt(8)
            style.font.italic = True
            style.font.color.rgb = RGBColor.from_string(MUTED)

    if "Source Code" not in styles:
        styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
    code = styles["Source Code"]
    code.font.name = "DejaVu Sans Mono"
    code.font.size = Pt(8)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Mm(5)
    code.paragraph_format.right_indent = Mm(5)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(5)

    doc.save(path)


def prepare_markdown(diagram: Path, path: Path) -> None:
    source = SOURCE.read_text(encoding="utf-8")
    marker = "\n---\n\n## 0."
    if marker not in source:
        raise RuntimeError("Unexpected source preamble")
    body = "## 0." + source.split(marker, 1)[1]
    # The canonical source reserves H1 for the document title. In the prepared
    # body the title comes from front matter, so promote every body heading once.
    body = re.sub(r"(?m)^(#{2,6})(?=\s)", lambda match: match.group(1)[1:], body)
    mermaid = re.compile(r"```mermaid\n.*?```", re.DOTALL)
    replacement = f'![Architektura MiLO: řízení práce je oddělené od zdroje pravdy.]({diagram.as_posix()}){{ width=16.8cm }}'
    body, count = mermaid.subn(replacement, body, count=1)
    if count != 1:
        raise RuntimeError("Architecture Mermaid block was not found exactly once")
    contents = """# Obsah

| Základ systému | Provoz a zavedení |
|---|---|
| 0. Výsledek v jedné větě | 11. Komunikace podle publika a cíle |
| 1. Co lze a nelze garantovat | 12. Integrace a sběr dat |
| 2. Pracovní profil Miroslava | 13. Okamžité MVP — přesný rozsah |
| 3. Jak vzniklo rozhodnutí | 14. Akceptační brána |
| 4. Architektura | 15. Pre-mortem |
| 5. Organizace agentů | 16. Provozní rytmus |
| 6. Čtyři pracovní workflow | 17. Metriky posunu |
| 7. Životní cyklus úkolu | 18. První provozní rozhodnutí |
| 8. Realtime monitoring | 19. Obsah MVP balíčku |
| 9. Kontrolní dashboard | 20. Zdroje a ověření nástrojů |
| 10. Znalostní báze | 21. Závěr |

"""
    front_matter = """---
title: "MiLO / Hermes"
subtitle: "Řídicí systém života a práce · návrh funkčního MVP"
author: "Pro Miroslava Brožka"
date: "21. července 2026 · verze 1.0"
lang: cs-CZ
toc-title: "Obsah"
---

"""
    path.write_text(front_matter + contents + body, encoding="utf-8")


def postprocess_docx(path: Path) -> None:
    doc = Document(path)
    doc.core_properties.title = "MiLO / Hermes — řídicí systém života a práce"
    doc.core_properties.subject = "Funkční MVP, agentní organizace, realtime monitoring a append-only zdroj pravdy"
    doc.core_properties.author = "MiLO design session pro Miroslava Brožka"
    doc.core_properties.keywords = "MiLO, Hermes, Paperclip, dashboard, agenti, Fact Ledger, TJ Krupka"

    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(19)
        section.right_margin = Mm(17)
        section.different_first_page_header_footer = True

        header = section.header
        hp = header.paragraphs[0]
        hp.text = "MiLO / Hermes  ·  Řídicí systém v1.0"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)

        footer = section.footer
        fp = footer.paragraphs[0]
        add_page_field(fp)

    # Make the cover a real first page and start the body after the contents.
    date_found = False
    body_started = False
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        if txt == "21. července 2026 · verze 1.0" and not date_found:
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            date_found = True
        if txt.startswith("0. Výsledek v jedné větě") and not body_started:
            paragraph.paragraph_format.page_break_before = True
            body_started = True

    for table in doc.tables:
        table.autofit = True
        if table.rows:
            set_repeat_table_header(table.rows[0])
            for cell in table.rows[0].cells:
                shade_cell(cell, NAVY)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)
                        run.font.size = Pt(8)
        for row_index, row in enumerate(table.rows[1:], start=1):
            if row_index % 2 == 0:
                for cell in row.cells:
                    shade_cell(cell, PALE)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(7.7)

    # Keep list items compact without flattening hierarchy.
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("List"):
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_together = True

    doc.save(path)


def main() -> None:
    BUILD.mkdir(parents=True, exist_ok=True)
    ARTIFACTS.mkdir(parents=True, exist_ok=True)
    diagram = build_architecture_diagram()
    reference = BUILD / "reference.docx"
    prepared = BUILD / "prepared.md"
    output = ARTIFACTS / "MiLO_Hermes_ridici_system_v1.docx"
    build_reference_doc(reference)
    prepare_markdown(diagram, prepared)
    subprocess.run(
        [
            "pandoc",
            str(prepared),
            "--from=markdown+pipe_tables+fenced_code_blocks+link_attributes",
            "--to=docx",
            f"--reference-doc={reference}",
            f"--resource-path={ROOT}",
            "--standalone",
            "-o",
            str(output),
        ],
        check=True,
    )
    postprocess_docx(output)
    print(output)


if __name__ == "__main__":
    main()
